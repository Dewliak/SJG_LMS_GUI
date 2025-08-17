import io
from typing import List

import qrcode
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

from datetime import datetime


from qr_book_model import QrBookModel

from lmsapi.context_api import Context

def chunks(lst, n):
    """Yield successive n-sized chunks from lst."""
    for i in range(0, len(lst), n):
        yield lst[i:i + n]

def make_text_printable(book_data_list: List[QrBookModel]) -> List[QrBookModel]:
    """
    Trims if necessary to make text printable
    """
    printable_book_data_list= []
    for element in book_data_list:
        element.name = element.name if len(element.name) <= 39 else element.name[:36] + "..."
        element.author = element.author if len(element.author) <= 39 else element.author[:36] + "..."

        printable_book_data_list.append(element)

    return printable_book_data_list


def set_cell_margins(table, left=0, right=0):
    tc = table._element
    tblPr = tc.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    kwargs = {"left": left, "right": right}
    for m in ["left", "right"]:
        node = OxmlElement("w:{}".format(m))
        node.set(qn('w:w'), str(kwargs.get(m)))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)

    tblPr.append(tblCellMar)

def generate_qr_code(link):
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=2.5,
        border=1,
    )
    qr.clear()
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save("qrcode.png")



def generate_qr_image(book_data_list: List[QrBookModel], context: Context, bytes: bool = False):
    doc = docx.Document()

    data = make_text_printable(book_data_list)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    table = doc.add_table(rows=0, cols=6)

    table.style = 'Table Grid'  # single lines in all cells
    table.autofit = False

    set_cell_margins(table)

    for i in range(3):
        table.columns[2 * i].width = Cm(2.5)
        table.columns[2 * i + 1].width = Cm(3.5)

    counter = 0

    for book_chunk in chunks(data, 3):

        counter = 0

        row = table.add_row().cells

        for book in book_chunk:
            paragraph = row[2 * counter].paragraphs[0]
            # --- add a run in which to place the picture ---
            run = paragraph.add_run()
            # --- add the picture to that run ---
            generate_qr_code(book.link)
            run.add_picture("qrcode.png")
            row[2 * counter + 1].text = book.id + "\n\n" + book.name + "\n\n" + book.author

            counter += 1

    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2)

    file_date = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")

    if bytes:
        bio = io.BytesIO()
        doc.save(bio)
        return doc, bio
    else:
        doc.save(f'generated_qr_images/könyv_matrica_{file_date}.docx')


if __name__ == "__main__":
    data = [QrBookModel("Matematika programozasok12312312312312313123123", "Obadovics J gyula",
                 "sulilms.vercel.app/?id=9a4b60e4-b44d-4cd0-94f6-ed0196091bf0"),
            QrBookModel("ASD", "1", 'http://127.0.0.1:5000/?=1'),
            QrBookModel("Matematika programozas 235", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 236", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 237", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 238", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 239", "Obadovics J gyula", "lindasasdasdasdasdasdata"),
            QrBookModel("Matematika programozas 2310", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 2311", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 2312", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 2313", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 2314", "Obadovics J gyula", "lindasdasdata"),
            QrBookModel("Matematika programozas 2315", "Obadovics J gyula", "lindasdasdata"),
            ]

    data2= [QrBookModel(name='ASD', author='1', link='http://127.0.0.1:5000/?=1'),
            QrBookModel(name='test', author='test', link='http://127.0.0.1:5000/?=562977095734'),]

    generate_qr_image(data,context=Context(), bytes=False)
